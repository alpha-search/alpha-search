"""Alpha Search \u2014 Real Data Strategy Research Report Generator.

Generates professional Markdown and DOCX reports from real-data backtest
results produced by the Alpha Search research pipelines.

Example::

    from alpha_search.research.report_writer import (
        generate_markdown_report, generate_docx_report, write_metadata,
        format_metrics_table,
    )

    md_path = generate_markdown_report(results, "/mnt/agents/output/report.md")
    docx_path = generate_docx_report(results, "/mnt/agents/output/report.docx")
    write_metadata(results, "/mnt/agents/output/metadata.json")
"""

from __future__ import annotations

import json
import logging
import os
from datetime import datetime, timezone
from typing import Any

import pandas as pd

logger = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────────────

_ALPHA_SEARCH_LAYERS: dict[str, str] = {
    "Data Layer": (
        "Fetches and normalises OHLCV data from Yahoo Finance (or other "
        "providers). Handles caching, missing-data imputation, and multi-ticker alignment."
    ),
    "Opportunity Discovery": (
        "Scans the universe for tradeable opportunities using momentum, "
        "mean-reversion, and statistical-arbitrage heuristics."
    ),
    "Signal Generation": (
        "Converts raw price data into actionable trading signals "
        "(SMA cross-overs, z-score thresholds, cointegration spreads)."
    ),
    "Backtest Engine": (
        "Event-driven backtester with transaction-cost modelling, "
        "slippage estimation, and portfolio-level P&L tracking."
    ),
    "Portfolio Construction": (
        "Builds optimal allocations via equal-weight, inverse-volatility, "
        "and risk-parity methods."
    ),
    "Memory Layer": (
        "Persistent agent memory (DuckDB/SQLite + Markdown journals) for "
        "strategy results, decisions, hand-offs, and risk flags."
    ),
    "Reporting": (
        "Produces human-readable Markdown / DOCX reports and JSON metadata "
        "for downstream analysis and archival."
    ),
}

_STRATEGIES: tuple[tuple[str, str, str], ...] = (
    ("momentum", "Momentum Strategy Results",
     "Trend-following strategy that exploits persistent price movements."),
    ("mean_reversion", "Mean Reversion Strategy Results",
     "Contrarian strategy that capitalises on price deviations from historical mean."),
    ("arbitrage", "Statistical Arbitrage Results",
     "Pairs-trading strategy that profits from mean-reverting spreads between cointegrated assets."),
)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _safe_serialize(obj: Any) -> Any:
    """Recursively convert non-JSON-serialisable objects to strings."""
    if isinstance(obj, pd.DataFrame):
        return obj.to_dict(orient="records")
    if isinstance(obj, pd.Series):
        return obj.to_dict()
    if isinstance(obj, (datetime, pd.Timestamp)):
        return obj.isoformat()
    if isinstance(obj, dict):
        return {k: _safe_serialize(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_safe_serialize(v) for v in obj]
    return obj


def _pct(value: float | None, decimals: int = 2) -> str:
    """Format a float as a percentage string, or ``'N/A'``."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "N/A"
    return f"{value * 100:.{decimals}f}%"


def _fmt(value: float | None, decimals: int = 4) -> str:
    """Format a float with given precision, or ``'N/A'``."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "N/A"
    return f"{value:.{decimals}f}"


def _df_to_md(df: pd.DataFrame | None, max_rows: int = 20) -> str:
    """Convert a DataFrame to a Markdown table (manual fallback if tabulate missing)."""
    if df is None or df.empty:
        return "*No data available.*"
    d = df.head(max_rows).copy()
    for col in d.columns:
        if pd.api.types.is_float_dtype(d[col]):
            d[col] = d[col].apply(lambda x: f"{x:.4f}" if pd.notna(x) else "N/A")
    try:
        return d.to_markdown(index=False)
    except Exception:
        cols = list(d.columns)
        lines = ["| " + " | ".join(str(c) for c in cols) + " |",
                 "| " + " | ".join("---" for _ in cols) + " |"]
        for _, row in d.iterrows():
            lines.append("| " + " | ".join(str(v) for v in row) + " |")
        return "\n".join(lines)


def _series_summary(s: pd.Series | None) -> str:
    """Render a short Markdown summary of a pandas Series."""
    if s is None or s.empty:
        return "*No series data available.*"
    return (f"- **Start:** {_fmt(s.iloc[0], 2)}\n"
            f"- **End:** {_fmt(s.iloc[-1], 2)}\n"
            f"- **Min:** {_fmt(s.min(), 2)}\n"
            f"- **Max:** {_fmt(s.max(), 2)}\n"
            f"- **Observations:** {len(s)}")


def _interpret(metrics: dict[str, Any]) -> str:
    """Generate a qualitative interpretation paragraph from metrics."""
    sharpe = metrics.get("sharpe_ratio") or metrics.get("sharpe")
    max_dd = metrics.get("max_drawdown")
    ann_ret = metrics.get("annualized_return")
    if sharpe is None or ann_ret is None or max_dd is None:
        return "*Insufficient metrics to provide a full interpretation.*"
    parts = []
    if sharpe > 1.0:
        parts.append(f"This strategy delivered a **strong Sharpe ratio of {sharpe:.2f}**,")
        parts.append("indicating attractive risk-adjusted returns.")
    elif sharpe > 0.5:
        parts.append(f"The strategy achieved a **moderate Sharpe ratio of {sharpe:.2f}**,")
        parts.append("suggesting some risk-adjusted alpha but with meaningful volatility.")
    else:
        parts.append(f"The Sharpe ratio of {sharpe:.2f} is **below typical thresholds**,")
        parts.append("indicating the strategy may not adequately compensate for risk taken.")
    parts.append(f"With an annualised return of {_pct(ann_ret)} and a maximum drawdown of")
    parts.append(f"{_pct(max_dd)}, the risk/reward profile should be carefully evaluated.")
    return " ".join(parts)


# ── Section builders (Markdown) ──────────────────────────────────────────────

def _section_disclaimer(results: dict[str, Any]) -> str:
    d = results.get("disclaimer", "RESEARCH / EDUCATIONAL PURPOSES ONLY. NOT INVESTMENT ADVICE.")
    return (f"## Disclaimer\n\n"
            f"> **IMPORTANT:** {d}\n\n"
            "This report presents quantitative backtest results for research and\n"
            "educational purposes. **Past performance does not guarantee future results.**\n"
            "These strategies have not been tested in live market conditions.\n\n---\n")


def _section_metadata(results: dict[str, Any]) -> str:
    m = results.get("metadata", {})
    tickers = m.get("tickers", [])
    td = ", ".join(tickers) if tickers else "N/A"
    if len(td) > 200:
        td = td[:200] + " ..."
    return (f"## Run Metadata\n\n| Parameter | Value |\n| --- | --- |\n"
            f"| Run Timestamp | {m.get('run_timestamp', 'N/A')} |\n"
            f"| Universe | {m.get('universe', 'N/A')} |\n"
            f"| Tickers | {m.get('n_tickers_fetched', len(tickers))} / "
            f"{m.get('n_tickers_requested', len(tickers))} fetched |\n"
            f"| Date Range | {m.get('start_date', 'N/A')} to {m.get('end_date', 'N/A')} |\n"
            f"| Data Source | {m.get('data_source', 'N/A')} |\n"
            f"| Initial Capital | {_fmt(m.get('initial_capital'), 0)} |\n"
            f"| Transaction Cost | {_pct(m.get('transaction_cost'), 4)} |\n\n"
            f"**Tickers:** {td}\n\n---\n")


def _section_architecture() -> str:
    lines = ["## Architecture Used\n\n",
             "This research run employed the following Alpha Search OS layers:\n"]
    for name, desc in _ALPHA_SEARCH_LAYERS.items():
        lines.append(f"**{name}** \u2014 {desc}\n")
    lines.append("---\n")
    return "\n".join(lines)


def _section_data_summary(results: dict[str, Any]) -> str:
    m = results.get("metadata", {})
    return (f"## Data Summary\n\n"
            f"- **Tickers Requested:** {m.get('n_tickers_requested', 'N/A')}\n"
            f"- **Tickers Fetched:** {m.get('n_tickers_fetched', 'N/A')}\n"
            f"- **Date Range:** {m.get('start_date', 'N/A')} to {m.get('end_date', 'N/A')}\n"
            f"- **Data Source:** {m.get('data_source', 'N/A')}\n"
            f"- **Initial Capital:** {_fmt(m.get('initial_capital'), 0)}\n"
            f"- **Transaction Cost:** {_pct(m.get('transaction_cost'), 4)}\n\n---\n")


def _section_liquidity(results: dict[str, Any]) -> str:
    liq = results.get("liquidity")
    table = _df_to_md(liq, max_rows=20) if liq is not None and not liq.empty else "*No liquidity data.*"
    return f"## Liquidity Summary\n\nOverview of trading liquidity across the universe:\n\n{table}\n\n---\n"


def _section_strategy(results: dict[str, Any], key: str, title: str, description: str) -> str:
    sd = results.get(key, {})
    metrics = sd.get("metrics", {})
    lines = [f"## {title}\n\n*({description})*\n\n",
             "### Performance Metrics\n\n", format_metrics_table(metrics), "\n\n",
             "### Interpretation\n\n", _interpret(metrics), "\n\n"]
    # Holdings / pairs table
    h_key = "pairs" if key == "arbitrage" else "holdings"
    h = sd.get(h_key)
    if h is not None and not h.empty:
        label = "Pair Information" if key == "arbitrage" else "Holdings Information"
        lines.extend([f"### {label}\n\n", _df_to_md(h, max_rows=15), "\n\n"])
    # Equity curve
    eq = sd.get("equity_curve")
    if eq is not None and not eq.empty:
        lines.extend(["### Equity Curve Summary\n\n", _series_summary(eq), "\n\n"])
    # Trade count
    trades = sd.get("trades")
    if trades is not None and not trades.empty:
        lines.append(f"- **Total Trades:** {len(trades)}\n\n")
    lines.append("---\n")
    return "\n".join(lines)


def _section_portfolio(results: dict[str, Any]) -> str:
    pf = results.get("portfolio", {})
    desc = {"equal_weight": "Equal capital allocation across all assets.",
            "inverse_vol": "Weights proportional to inverse realised volatility.",
            "risk_parity": "Risk-budget allocation targeting equal risk contribution."}
    lines = ["## Portfolio Optimization Results\n\n",
             "Comparison of portfolio construction methodologies:\n\n",
             "| Method | Description | Sharpe | Ann. Return | Max DD |\n",
             "| --- | --- | --- | --- | --- |\n"]
    for mk in ("equal_weight", "inverse_vol", "risk_parity"):
        md = pf.get(mk, {})
        if not md:
            continue
        m = md.get("metrics", {})
        lines.append(f"| {mk.replace('_', ' ').title()} | {desc.get(mk, '')} | "
                     f"{_fmt(m.get('sharpe_ratio') or m.get('sharpe'), 2)} | "
                     f"{_pct(m.get('annualized_return'))} | {_pct(m.get('max_drawdown'))} |\n")
    lines.append("\n---\n")
    return "".join(lines)


def _section_sharpe_drawdown(results: dict[str, Any]) -> str:
    lines = ["## Sharpe Ratio & Drawdown Summary\n\n",
             "| Strategy | Sharpe Ratio | Max Drawdown | Annualized Return |\n",
             "| --- | --- | --- | --- |\n"]
    best_key, best_val = None, -float("inf")
    for key, title, _ in _STRATEGIES:
        m = results.get(key, {}).get("metrics", {})
        s = m.get("sharpe_ratio") or m.get("sharpe")
        if s is not None and s > best_val:
            best_val, best_key = s, key
        lines.append(f"| {title.split(' Results')[0]} | {_fmt(s, 2)} | "
                     f"{_pct(m.get('max_drawdown'))} | {_pct(m.get('annualized_return'))} |\n")
    pf = results.get("portfolio", {})
    for mk in ("equal_weight", "inverse_vol", "risk_parity"):
        md = pf.get(mk, {})
        if not md:
            continue
        m = md.get("metrics", {})
        s = m.get("sharpe_ratio") or m.get("sharpe")
        if s is not None and s > best_val:
            best_val, best_key = s, f"pf_{mk}"
        lines.append(f"| Portfolio ({mk.replace('_', ' ').title()}) | {_fmt(s, 2)} | "
                     f"{_pct(m.get('max_drawdown'))} | {_pct(m.get('annualized_return'))} |\n")
    lines.append("\n")
    if best_key:
        if best_key.startswith("pf_"):
            name = best_key[3:].replace("_", " ").title()
            lines.append(f"**Best performer:** Portfolio ({name}) with Sharpe **{best_val:.2f}**.\n\n")
        else:
            for key, title, _ in _STRATEGIES:
                if key == best_key:
                    lines.append(f"**Best performer:** {title.split(' Results')[0]} with Sharpe "
                                 f"**{best_val:.2f}**.\n\n")
                    break
        lines.append("The highest Sharpe ratio indicates the best risk-adjusted returns. "
                     "Consider drawdown, costs, and correlation before combining strategies.\n")
    else:
        lines.append("*Sharpe ratios not available for comparison.*\n")
    lines.append("\n---\n")
    return "".join(lines)


def _section_risks() -> str:
    items = [
        "**Data Quality**: Yahoo Finance data may contain unadjusted splits, dividends, and missing points.",
        "**Transaction Costs**: Flat 10 bps per trade; real-world costs may be higher for high-turnover strategies.",
        "**Slippage**: No slippage model applied; execution at quoted prices is not guaranteed.",
        "**Overfitting**: Parameters selected in-sample; walk-forward validation required for robustness.",
        "**Survivorship Bias**: Only currently-listed large-cap equities are included.",
        "**Yahoo Finance Limitations**: Delayed quotes, limited history, rate limiting on free tier.",
        "**No Live Execution**: Results are purely simulated; market microstructure can alter live performance.",
    ]
    return "## Risks and Limitations\n\n" + "\n\n".join(f"{i+1}. {item}" for i, item in enumerate(items)) + "\n\n---\n"


def _section_next_actions() -> str:
    items = [
        "**Walk-Forward Validation**: Rolling-window cross-validation across multiple market regimes.",
        "**Intraday Data**: 1-minute or tick-level data for reduced signal latency.",
        "**Advanced Cost Model**: Market-impact models and volume-weighted slippage.",
        "**Larger Universe**: Mid-caps, international markets, and sector ETFs.",
        "**Sector-Neutral Construction**: Equal sector exposure to eliminate unintended sector bets.",
        "**Crypto / India Expansion**: Cross-market generalisability assessment.",
        "**Regime Detection**: Dynamic parameter adaptation based on market regime classification.",
    ]
    return "## Next Research Actions\n\n" + "\n\n".join(f"{i+1}. {item}" for i, item in enumerate(items)) + "\n\n---\n"


def _section_appendix(results: dict[str, Any]) -> str:
    lines = ["## Appendix: Memory Records Summary\n\n",
             "Summary of agent memory records associated with this research run:\n\n"]
    records = results.get("memory_records", [])
    if records:
        lines.append("| Agent | Type | Title | Status | Importance |\n| --- | --- | --- | --- | --- |\n")
        for r in records:
            lines.append(f"| {r.get('agent_name','N/A')} | {r.get('memory_type','N/A')} | "
                         f"{r.get('title','N/A')} | {r.get('status','N/A')} | "
                         f"{r.get('importance_score','N/A')} |\n")
        lines.append("\n")
    else:
        lines.append("*No memory records embedded. Check the Alpha Search memory store (DuckDB/SQLite) "
                     "for persisted strategy results, hand-offs, and risk decisions.*\n\n")
    sm = results.get("strategy_memories", [])
    if sm:
        lines.append("### Strategy Memory Entries\n\n")
        for mem in sm:
            lines.append(f"- **[{mem.get('verdict','unknown')}]** {mem.get('strategy_name','Unknown')} "
                         f"\u2014 Sharpe: {_fmt(mem.get('sharpe'), 2)}, Max DD: {_pct(mem.get('max_drawdown'))}\n")
    lines.append("\n---\n")
    return "".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

def format_metrics_table(metrics: dict[str, Any]) -> str:
    """Format a metrics dictionary as a Markdown table.

    Args:
        metrics: Dictionary mapping metric name to numeric value.

    Returns:
        Markdown table string (two columns: Metric, Value).
    """
    if not metrics:
        return "*No metrics available.*"
    lines = ["| Metric | Value |", "| --- | --- |"]
    for key, value in sorted(metrics.items()):
        label = key.replace("_", " ").title()
        if any(w in key.lower() for w in ("return", "drawdown")):
            formatted = _pct(value)
        elif any(w in key.lower() for w in ("ratio", "rate")):
            formatted = _fmt(value, 4)
        elif "cost" in key.lower():
            formatted = _pct(value, 4)
        else:
            formatted = _fmt(value, 4)
        lines.append(f"| {label} | {formatted} |")
    return "\n".join(lines)


def generate_markdown_report(results: dict[str, Any], output_path: str) -> str:
    """Generate a comprehensive Markdown report from backtest results.

    Args:
        results: Nested results dict with keys ``metadata``, ``momentum``,
            ``mean_reversion``, ``arbitrage``, ``portfolio``, ``liquidity``,
            and optional ``disclaimer``, ``memory_records``.
        output_path: Destination file path (``.md`` extension).

    Returns:
        Absolute path to the written file.

    Raises:
        OSError: If the file cannot be written.
    """
    if not results:
        logger.warning("Empty results dict passed to generate_markdown_report.")

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    header = (f"# Alpha Search \u2014 Real Data Strategy Research Report\n\n"
              f"*Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}*\n\n---\n")
    sections = [
        header,
        _section_disclaimer(results),
        _section_metadata(results),
        _section_architecture(),
        _section_data_summary(results),
        _section_liquidity(results),
    ]
    for key, title, desc in _STRATEGIES:
        sections.append(_section_strategy(results, key, title, desc))
    sections.extend([
        _section_portfolio(results),
        _section_sharpe_drawdown(results),
        _section_risks(),
        _section_next_actions(),
        _section_appendix(results),
    ])

    text = "\n".join(sections)
    try:
        with open(output_path, "w", encoding="utf-8") as fh:
            fh.write(text)
        logger.info("Markdown report written to %s", output_path)
    except OSError as exc:
        logger.error("Failed to write Markdown report to %s: %s", output_path, exc)
        raise
    return os.path.abspath(output_path)


def generate_docx_report(results: dict[str, Any], output_path: str) -> str:
    """Generate a DOCX report from backtest results.

    Uses ``python-docx`` when available; otherwise writes a structured
    text file with ``.docx`` extension and logs a warning.

    Args:
        results: Nested results dict (same structure as Markdown report).
        output_path: Destination file path (``.docx`` extension).

    Returns:
        Absolute path to the written file.

    Raises:
        OSError: If the file cannot be written.
    """
    if not results:
        logger.warning("Empty results dict passed to generate_docx_report.")

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

        doc = Document()

        # Title
        title = doc.add_heading("Alpha Search \u2014 Real Data Strategy Research Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}").italic = True

        # Disclaimer
        doc.add_heading("Disclaimer", level=1)
        d = results.get("disclaimer", "RESEARCH / EDUCATIONAL PURPOSES ONLY. NOT INVESTMENT ADVICE.")
        doc.add_paragraph(f"IMPORTANT: {d} Past performance does not guarantee future results.")

        # Metadata
        doc.add_heading("Run Metadata", level=1)
        m = results.get("metadata", {})
        _docx_table(doc, [("Parameter", "Value")] + [
            ("Run Timestamp", m.get("run_timestamp", "N/A")),
            ("Universe", m.get("universe", "N/A")),
            ("Tickers", f"{m.get('n_tickers_fetched','N/A')} / {m.get('n_tickers_requested','N/A')} fetched"),
            ("Date Range", f"{m.get('start_date','N/A')} to {m.get('end_date','N/A')}"),
            ("Data Source", m.get("data_source", "N/A")),
            ("Initial Capital", _fmt(m.get("initial_capital"), 0)),
            ("Transaction Cost", _pct(m.get("transaction_cost"), 4)),
        ])

        # Architecture
        doc.add_heading("Architecture Used", level=1)
        for name, desc in _ALPHA_SEARCH_LAYERS.items():
            p = doc.add_paragraph()
            p.add_run(f"{name} ").bold = True
            p.add_run(f"\u2014 {desc}")

        # Data Summary
        doc.add_heading("Data Summary", level=1)
        m = results.get("metadata", {})
        for label, val in [("Tickers Requested", m.get('n_tickers_requested','N/A')),
                           ("Tickers Fetched", m.get('n_tickers_fetched','N/A')),
                           ("Date Range", f"{m.get('start_date','N/A')} to {m.get('end_date','N/A')}"),
                           ("Data Source", m.get('data_source','N/A')),
                           ("Initial Capital", _fmt(m.get('initial_capital'), 0)),
                           ("Transaction Cost", _pct(m.get('transaction_cost'), 4))]:
            doc.add_paragraph(f"{label}: {val}", style="List Bullet")

        # Liquidity
        doc.add_heading("Liquidity Summary", level=1)
        liq = results.get("liquidity")
        if liq is not None and not liq.empty:
            _docx_df(doc, liq.head(20))
        else:
            doc.add_paragraph("No liquidity data available.")

        # Strategy sections
        for key, title, description in _STRATEGIES:
            doc.add_heading(title, level=1)
            doc.add_paragraph(description, style="Intense Quote")
            sd = results.get(key, {})
            metrics = sd.get("metrics", {})
            doc.add_heading("Performance Metrics", level=2)
            _docx_metrics(doc, metrics)
            h_key = "pairs" if key == "arbitrage" else "holdings"
            h = sd.get(h_key)
            if h is not None and not h.empty:
                doc.add_heading("Pair Information" if key == "arbitrage" else "Holdings Information", level=2)
                _docx_df(doc, h.head(15))

        # Portfolio
        doc.add_heading("Portfolio Optimization Results", level=1)
        pf = results.get("portfolio", {})
        rows = [("Method", "Sharpe", "Ann. Return", "Max DD")]
        for mk in ("equal_weight", "inverse_vol", "risk_parity"):
            md = pf.get(mk, {})
            if md:
                m = md.get("metrics", {})
                rows.append((mk.replace("_", " ").title(),
                             _fmt(m.get("sharpe_ratio") or m.get("sharpe"), 2),
                             _pct(m.get("annualized_return")),
                             _pct(m.get("max_drawdown"))))
        _docx_table(doc, rows)

        # Sharpe summary
        doc.add_heading("Sharpe Ratio & Drawdown Summary", level=1)
        rows = [("Strategy", "Sharpe Ratio", "Max Drawdown", "Ann. Return")]
        for key, title, _ in _STRATEGIES:
            m = results.get(key, {}).get("metrics", {})
            rows.append((title.split(" Results")[0],
                         _fmt(m.get("sharpe_ratio") or m.get("sharpe"), 2),
                         _pct(m.get("max_drawdown")),
                         _pct(m.get("annualized_return"))))
        _docx_table(doc, rows)

        # Risks
        doc.add_heading("Risks and Limitations", level=1)
        for item in [
            "Data Quality: Yahoo Finance data may contain unadjusted splits, dividends, and missing points.",
            "Transaction Costs: Flat 10 bps per trade; real-world costs may be higher.",
            "Slippage: No slippage model applied; execution at quoted prices not guaranteed.",
            "Overfitting: Parameters selected in-sample; walk-forward validation required.",
            "Survivorship Bias: Only currently-listed large-cap equities included.",
            "Yahoo Finance Limitations: Delayed quotes, limited history, rate limiting.",
            "No Live Execution: Results are purely simulated.",
        ]:
            doc.add_paragraph(item, style="List Number")

        # Next actions
        doc.add_heading("Next Research Actions", level=1)
        for item in [
            "Walk-Forward Validation: Rolling-window cross-validation across market regimes.",
            "Intraday Data: 1-minute or tick-level data for reduced signal latency.",
            "Advanced Cost Model: Market-impact models and volume-weighted slippage.",
            "Larger Universe: Mid-caps, international markets, sector ETFs.",
            "Sector-Neutral Construction: Equal sector exposure to eliminate sector bets.",
            "Crypto / India Expansion: Cross-market generalisability assessment.",
            "Regime Detection: Dynamic parameter adaptation based on market regime.",
        ]:
            doc.add_paragraph(item, style="List Number")

        # Appendix
        doc.add_heading("Appendix: Memory Records Summary", level=1)
        records = results.get("memory_records", [])
        if records:
            rows = [("Agent", "Type", "Title", "Status", "Importance")]
            for r in records:
                rows.append((str(r.get("agent_name","N/A")), str(r.get("memory_type","N/A")),
                             str(r.get("title","N/A")), str(r.get("status","N/A")),
                             str(r.get("importance_score","N/A"))))
            _docx_table(doc, rows)
        else:
            doc.add_paragraph("No memory records embedded. Check the Alpha Search memory store.")

        doc.save(output_path)
        logger.info("DOCX report written to %s", output_path)

    except ImportError:
        logger.warning("python-docx not installed. Writing structured text fallback with .docx extension.")
        md_path = output_path.replace(".docx", "_fallback.md")
        try:
            generate_markdown_report(results, md_path)
            with open(md_path, "r", encoding="utf-8") as fh:
                content = fh.read()
        except Exception as exc:
            logger.error("Failed to generate fallback content: %s", exc)
            content = "Error generating report content."
        header = ("ALPHA SEARCH \u2014 REAL DATA STRATEGY RESEARCH REPORT\n"
                  "=" * 72 + "\n\n"
                  f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n\n"
                  "NOTE: Structured text fallback \u2014 python-docx is not installed.\n\n")
        with open(output_path, "w", encoding="utf-8") as fh:
            fh.write(header + content)
        logger.info("DOCX fallback written to %s", output_path)

    except OSError as exc:
        logger.error("Failed to write DOCX report to %s: %s", output_path, exc)
        raise

    return os.path.abspath(output_path)


# ── DOCX helper functions ────────────────────────────────────────────────────

def _docx_table(doc: Any, rows: list[tuple[Any, ...]]) -> None:
    """Add a table to a python-docx Document from a list of row tuples.

    Args:
        doc: A ``Document`` instance.
        rows: List of tuples, where the first tuple is the header.
    """
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for idx, val in enumerate(rows[0]):
        table.rows[0].cells[idx].text = str(val)
    for row_data in rows[1:]:
        cells = table.add_row().cells
        for idx, val in enumerate(row_data):
            cells[idx].text = str(val)


def _docx_df(doc: Any, df: pd.DataFrame) -> None:
    """Add a DataFrame as a table to a python-docx Document.

    Args:
        doc: A ``Document`` instance.
        df: DataFrame to render.
    """
    if df.empty:
        doc.add_paragraph("(empty table)")
        return
    rows: list[tuple[Any, ...]] = [tuple(str(c) for c in df.columns)]
    for _, row in df.iterrows():
        rows.append(tuple(f"{v:.4f}" if isinstance(v, float) else str(v) for v in row))
    _docx_table(doc, rows)


def _docx_metrics(doc: Any, metrics: dict[str, Any]) -> None:
    """Add a metrics dictionary as a two-column table to a python-docx Document.

    Args:
        doc: A ``Document`` instance.
        metrics: Dictionary of metric name -> value.
    """
    if not metrics:
        doc.add_paragraph("No metrics available.")
        return
    rows: list[tuple[Any, ...]] = [("Metric", "Value")]
    for key, value in sorted(metrics.items()):
        label = key.replace("_", " ").title()
        if any(w in key.lower() for w in ("return", "drawdown")):
            formatted = _pct(value)
        elif any(w in key.lower() for w in ("ratio", "rate")):
            formatted = _fmt(value, 4)
        elif "cost" in key.lower():
            formatted = _pct(value, 4)
        else:
            formatted = _fmt(value, 4)
        rows.append((label, formatted))
    _docx_table(doc, rows)


# ── Metadata writer ──────────────────────────────────────────────────────────

def write_metadata(results: dict[str, Any], output_path: str) -> None:
    """Write a JSON metadata file with all run parameters.

    Non-serialisable objects (DataFrames, Series, datetime) are converted
    to JSON-friendly representations via :func:`_safe_serialize`.

    Args:
        results: Full results dictionary from the research pipeline.
        output_path: Destination file path (``.json`` extension).

    Raises:
        OSError: If the file cannot be written.
        TypeError: If the metadata cannot be JSON-serialised.
    """
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    meta = _safe_serialize(results.get("metadata", {}))
    meta["report_generated_at"] = datetime.now(timezone.utc).isoformat()
    meta["strategies_evaluated"] = [k for k in ("momentum", "mean_reversion", "arbitrage") if k in results]
    meta["portfolio_methods"] = list(results.get("portfolio", {}).keys())

    try:
        with open(output_path, "w", encoding="utf-8") as fh:
            json.dump(meta, fh, indent=2, ensure_ascii=False)
        logger.info("Metadata written to %s", output_path)
    except (OSError, TypeError) as exc:
        logger.error("Failed to write metadata to %s: %s", output_path, exc)
        raise
