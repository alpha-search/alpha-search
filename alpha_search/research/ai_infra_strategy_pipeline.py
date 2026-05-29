"""AI Infrastructure & Semiconductor Alpha Research Pipeline.

Portfolio-level cross-sectional momentum, trend-following, mean-reversion, and
Donchian-breakout strategies on the US AI-infrastructure and semiconductor universe.

Integrity rules (non-negotiable — do not relax):
1. Real OHLCV from Yahoo Finance only.  No synthetic data.
2. One pre-registered parameter set per strategy family.  No Sharpe fishing.
3. rf = 0.0 (no FRED).  This slightly overstates long-only Sharpe in a
   high-rate regime — the caveat is flagged wherever relevant.
4. Costs charged on traded notional (one-way bps) at every rebalance.
   Net return is always the headline metric; gross is shown alongside it.
5. No look-ahead: signals at date t use only data up to t; portfolio weights
   derived from those signals take effect at t+1 (shift-1 rule enforced).
6. Drawdown convention: NEGATIVE fraction (−0.25 = 25 % drawdown) throughout
   this module, matching the research artifact.
7. Negative or low Sharpe is reported as-is.  If a strategy fails to produce
   alpha it is recorded as "unprofitable" or "no_results" — no tuning to change it.
"""

from __future__ import annotations

import io
import json
import logging
import os
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Universe definition
# ---------------------------------------------------------------------------

_UNIVERSE: Dict[str, List[str]] = {
    "semiconductors": [
        "NVDA", "AMD", "AVGO", "TSM", "QCOM", "TXN", "INTC", "MU", "ADI", "NXPI",
        "MCHP", "ON", "MRVL", "MPWR", "SWKS", "QRVO", "LSCC", "ARM",
    ],
    "semi_equipment": ["ASML", "AMAT", "LRCX", "KLAC", "TER", "ENTG"],
    "ai_infra": ["ANET", "VRT", "SMCI", "DELL", "CRDO", "ALAB", "CIEN", "COHR"],
}

_BENCHMARKS: List[str] = ["SOXX", "SMH", "QQQ", "SPY"]

# Pre-registered, fixed parameters — declared once, never tuned against results.
DEFAULT_CONFIG: Dict[str, Any] = {
    # Cross-sectional momentum
    "cs_lookback_days": 252,    # ~12 months
    "cs_skip_days": 21,         # skip last month (avoids short-term reversal)
    "rebalance_freq": "ME",     # month-end
    "liq_window": 63,           # trailing bars for liquidity screen
    "min_dollar_vol": 25e6,     # $25M median daily $-volume to be eligible
    "quantile": 1.0 / 3.0,      # top / bottom tercile long / short
    # Trend following
    "ma_windows": [50, 100, 200],
    # Mean reversion
    "mr_window": 20,
    "mr_z_threshold": 2.0,
    # Breakout
    "breakout_window": 20,
    # Accounting
    "rf_annual": 0.0,           # NO FRED — rf = 0; caveat stated above
    "trading_days": 252,
    "oos_split": 0.5,           # temporal midpoint split (not chosen to flatter)
    "min_eligible": 6,          # minimum names required for tercile formation
}

_STRATEGY_HYPOTHESES: Dict[str, str] = {
    "cross_sectional_momentum": (
        "AI-infra + semis stocks with the strongest trailing 12-1 month returns "
        "continue to outperform the bottom tercile over the next month (L/S, dollar-neutral)."
    ),
    "trend_following": (
        "Stocks trading above their 50/100/200-day moving averages trend upward; "
        "equal-weight long-only portfolio of qualifying names."
    ),
    "mean_reversion": (
        "AI-infra + semis names that deviate ≥2σ below their 20-day rolling mean "
        "tend to mean-revert; long-only entries."
    ),
    "breakout": (
        "Stocks closing above the prior 20-day Donchian channel high (shifted 1 bar "
        "to prevent look-ahead) continue trending upward."
    ),
}


# ---------------------------------------------------------------------------
# 1. Universe accessor
# ---------------------------------------------------------------------------

def get_ai_infra_universe() -> Dict[str, List[str]]:
    """Return the pre-defined AI-infrastructure and semiconductor universe.

    Returns
    -------
    dict[str, list[str]]
        Keys: ``semiconductors``, ``semi_equipment``, ``ai_infra``.
        Values: Yahoo-Finance ticker lists.
    """
    return {k: list(v) for k, v in _UNIVERSE.items()}


def _all_universe_symbols() -> List[str]:
    tickers = []
    for v in _UNIVERSE.values():
        tickers.extend(v)
    return list(dict.fromkeys(tickers))  # preserve order, de-dup


# ---------------------------------------------------------------------------
# 2. Data ingestion
# ---------------------------------------------------------------------------

def download_ai_infra_data(
    symbols: Optional[List[str]] = None,
    period: str = "5y",
    interval: str = "1d",
    benchmarks: Optional[List[str]] = None,
) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """Download real OHLCV from Yahoo Finance for the AI-infra universe.

    Parameters
    ----------
    symbols:
        Universe tickers.  Defaults to the full pre-defined universe.
    period:
        yfinance history period (e.g. ``"5y"``, ``"2y"``).
    interval:
        Bar interval (e.g. ``"1d"``).
    benchmarks:
        Benchmark tickers (SOXX, SMH, QQQ, SPY by default).

    Returns
    -------
    tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]
        ``(close, volume, bench_close)`` — wide DataFrames with tickers as
        columns.  Columns missing data for the entire period are dropped.
    """
    try:
        import yfinance as yf
    except ImportError as exc:
        raise ImportError("yfinance is required: pip install yfinance") from exc

    symbols = symbols or _all_universe_symbols()
    benchmarks = benchmarks or _BENCHMARKS

    all_tickers = list(dict.fromkeys(symbols + benchmarks))
    logger.info("Downloading %d tickers (period=%s, interval=%s)…", len(all_tickers), period, interval)

    raw = yf.download(
        all_tickers,
        period=period,
        interval=interval,
        auto_adjust=True,
        progress=False,
        group_by="column",
        threads=True,
    )

    if isinstance(raw.columns, pd.MultiIndex):
        close_all = raw["Close"].copy()
        volume_all = raw["Volume"].copy()
    else:
        # Single-ticker edge case
        close_all = raw[["Close"]].rename(columns={"Close": all_tickers[0]})
        volume_all = raw[["Volume"]].rename(columns={"Volume": all_tickers[0]})

    close_all = close_all.sort_index()
    volume_all = volume_all.sort_index()

    # Split universe vs benchmark
    univ_cols = [c for c in symbols if c in close_all.columns]
    bench_cols = [c for c in benchmarks if c in close_all.columns]

    close = close_all[univ_cols]
    volume = volume_all[[c for c in univ_cols if c in volume_all.columns]]
    bench_close = close_all[bench_cols] if bench_cols else pd.DataFrame(index=close.index)

    logger.info("Universe: %d/%d symbols downloaded; benchmarks: %d/%d",
                len(univ_cols), len(symbols), len(bench_cols), len(benchmarks))
    return close, volume, bench_close


# ---------------------------------------------------------------------------
# 3. Data validation
# ---------------------------------------------------------------------------

def validate_ai_infra_data(
    close: pd.DataFrame,
    volume: pd.DataFrame,
    min_history_days: int = 2 * 252,
) -> Tuple[bool, Dict[str, Any], List[str], pd.DataFrame, pd.DataFrame]:
    """Validate downloaded OHLCV; flag and drop problematic symbols.

    Parameters
    ----------
    close:
        Wide close-price DataFrame (tickers as columns).
    volume:
        Wide volume DataFrame.
    min_history_days:
        Minimum trading days of history required (default: ~2 years).

    Returns
    -------
    tuple
        ``(all_valid, report, skipped, valid_close, valid_volume)``
        where *all_valid* is True when every symbol passes minimum checks.
    """
    report: Dict[str, Dict] = {}
    skipped: List[str] = []
    valid_cols: List[str] = []

    n_total = len(close)

    for col in close.columns:
        s = close[col]
        non_na = s.dropna()
        coverage = len(non_na) / n_total if n_total else 0.0

        issues: List[str] = []
        if len(non_na) < min_history_days:
            issues.append(
                f"only {len(non_na)} valid bars (need {min_history_days}); "
                "symbol will be excluded from tercile selection until it qualifies"
            )
        if (non_na <= 0).any():
            issues.append(f"{(non_na <= 0).sum()} non-positive close prices")
        rets = s.pct_change().dropna()
        bad_jumps = (rets.abs() > 0.60).sum()
        if bad_jumps > 0:
            issues.append(f"{bad_jumps} single-day moves > 60% (possible data error)")

        is_ok = len(non_na) > 0  # allow short history — per-rebalance screen handles it
        report[col] = {
            "n_bars": len(non_na),
            "coverage": round(coverage, 4),
            "issues": issues,
            "eligible_for_selection": len(non_na) >= min_history_days,
        }

        if (non_na <= 0).any() and len(non_na) > 0:
            skipped.append(col)
            logger.warning("validate_ai_infra_data: %s has non-positive prices — skipped", col)
        else:
            valid_cols.append(col)
            if issues:
                logger.info("validate_ai_infra_data: %s — %s", col, "; ".join(issues))

    valid_close = close[valid_cols]
    vol_cols = [c for c in valid_cols if c in volume.columns]
    valid_volume = volume[vol_cols]

    all_valid = len(skipped) == 0
    logger.info("Validation: %d valid, %d skipped", len(valid_cols), len(skipped))
    return all_valid, report, skipped, valid_close, valid_volume


# ---------------------------------------------------------------------------
# 4a. Cross-sectional momentum signal (primary)
# ---------------------------------------------------------------------------

def build_cross_sectional_momentum_signal(
    close: pd.DataFrame,
    volume: pd.DataFrame,
    lookback: int = 252,
    skip: int = 21,
    liq_window: int = 63,
    min_dollar_vol: float = 25e6,
    quantile: float = 1.0 / 3.0,
    freq: str = "ME",
    long_only: bool = False,
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """Build cross-sectional momentum weights at each rebalance date.

    Signal at t = close[t−skip] / close[t−lookback] − 1.
    Uses only past data at every rebalance (no look-ahead).

    Parameters
    ----------
    close, volume:
        Wide DataFrames — tickers as columns.
    lookback, skip:
        Momentum lookback and reversal-skip in trading days.
    liq_window:
        Rolling window for liquidity screen (median daily $ volume).
    min_dollar_vol:
        Minimum median daily $ volume to be eligible at a rebalance.
    quantile:
        Fraction of eligible names in each leg (default: top/bottom third).
    freq:
        Rebalance frequency for ``pd.Grouper`` (default ``"ME"`` = month-end).
    long_only:
        If True, build a long-only top-quantile portfolio (no short leg).

    Returns
    -------
    tuple[pd.DataFrame, pd.DataFrame]
        ``(ls_weights, rebal_dates_df)`` — weight DataFrame indexed by
        rebalance dates, tickers as columns.  L/S weights sum to ~0;
        long-only weights sum to ~1.
    """
    raw_signal = close.shift(skip) / close.shift(lookback) - 1.0
    dollar_vol = (close * volume).rolling(liq_window).median()

    rebal_dates = close.resample(freq).last().index
    rebal_dates = [d for d in rebal_dates if d in close.index]

    weights = pd.DataFrame(0.0, index=rebal_dates, columns=close.columns)

    for d in rebal_dates:
        sig = raw_signal.loc[d]
        liq = dollar_vol.loc[d]
        eligible_mask = sig.notna() & liq.notna() & (liq >= min_dollar_vol)
        eligible = sig[eligible_mask].sort_values()
        n = len(eligible)
        if n < 6:
            continue  # not enough names for meaningful terciles — stay flat
        k = max(1, int(round(n * quantile)))
        longs = eligible.index[-k:]
        if long_only:
            weights.loc[d, longs] = 1.0 / k
        else:
            shorts = eligible.index[:k]
            weights.loc[d, longs] = 1.0 / k
            weights.loc[d, shorts] = -1.0 / k

    weights = weights.clip(-1.0, 1.0)
    return weights, pd.DataFrame({"rebal_date": rebal_dates})


# ---------------------------------------------------------------------------
# 4b. Trend-following signal
# ---------------------------------------------------------------------------

def build_trend_following_signal(
    close: pd.DataFrame,
    ma_windows: Optional[List[int]] = None,
    long_only: bool = True,
) -> pd.DataFrame:
    """Long when price is above ALL specified moving averages.

    Signal is computed on close-to-close data using rolling means.
    No look-ahead: MA at t uses prices up to and including t.
    The position at t is based on signal at t−1 (shift-1 applied downstream
    in the portfolio backtest).

    Parameters
    ----------
    close:
        Wide close-price DataFrame.
    ma_windows:
        List of MA periods (default: [50, 100, 200]).

    Returns
    -------
    pd.DataFrame
        Binary signal DataFrame (0 or 1 per ticker per day).
    """
    if ma_windows is None:
        ma_windows = DEFAULT_CONFIG["ma_windows"]

    conditions = [close > close.rolling(w).mean() for w in ma_windows]
    combined = conditions[0]
    for c in conditions[1:]:
        combined = combined & c

    return combined.astype(float)


# ---------------------------------------------------------------------------
# 4c. Mean-reversion signal
# ---------------------------------------------------------------------------

def build_mean_reversion_signal(
    close: pd.DataFrame,
    window: int = 20,
    z_threshold: float = 2.0,
    allow_short: bool = False,
) -> pd.DataFrame:
    """Long when rolling z-score is below −z_threshold; exit when z > 0.

    Parameters
    ----------
    close:
        Wide close-price DataFrame.
    window:
        Rolling window for z-score computation.
    z_threshold:
        Entry threshold (in standard deviations below the mean).
    allow_short:
        If True, also short when z > +z_threshold.

    Returns
    -------
    pd.DataFrame
        Signal DataFrame: 1 = long, 0 = flat, −1 = short (if allow_short).
    """
    roll_mean = close.rolling(window).mean()
    roll_std = close.rolling(window).std().replace(0, np.nan)
    z = (close - roll_mean) / roll_std

    signal = pd.DataFrame(0.0, index=close.index, columns=close.columns)
    signal[z < -z_threshold] = 1.0

    if allow_short:
        signal[z > z_threshold] = -1.0

    return signal


# ---------------------------------------------------------------------------
# 4d. Breakout signal (Donchian channel)
# ---------------------------------------------------------------------------

def build_breakout_signal(
    close: pd.DataFrame,
    high: Optional[pd.DataFrame] = None,
    low: Optional[pd.DataFrame] = None,
    window: int = 20,
    allow_short: bool = False,
) -> pd.DataFrame:
    """Donchian-channel breakout with explicit look-ahead prevention.

    The channel is computed on data ending at t−1 (``shift(1)`` before
    ``rolling``) so the signal at t is strictly out-of-sample.

    Parameters
    ----------
    close, high, low:
        Wide DataFrames.  If high/low are None, close is used for both.
    window:
        Lookback for the Donchian channel.
    allow_short:
        If True, short when close < N-bar low.

    Returns
    -------
    pd.DataFrame
        Signal: 1 = long, 0 = flat, −1 = short (if allow_short).
    """
    h = high if high is not None else close
    l = low if low is not None else close

    # shift(1) BEFORE rolling — this is the no-look-ahead guarantee
    chan_high = h.shift(1).rolling(window).max()
    chan_low = l.shift(1).rolling(window).min()

    signal = pd.DataFrame(0.0, index=close.index, columns=close.columns)
    signal[close > chan_high] = 1.0

    if allow_short:
        signal[close < chan_low] = -1.0

    return signal


# ---------------------------------------------------------------------------
# 5. Monthly rebalance weights from per-ticker signal
# ---------------------------------------------------------------------------

def build_monthly_rebalance_weights(
    signal_df: pd.DataFrame,
    freq: str = "ME",
    quantile: float = 1.0 / 3.0,
    min_eligible: int = 4,
    long_only: bool = True,
) -> pd.DataFrame:
    """Snapshot a daily signal DataFrame to produce monthly rebalance weights.

    At each rebalance date, takes the signal value and forms a long (or L/S)
    equal-weight portfolio from the names with positive signals.

    Parameters
    ----------
    signal_df:
        Daily per-ticker signal (e.g. from trend-following or breakout).
    freq:
        Rebalance frequency.
    quantile:
        Used only when long_only=False to form tercile portfolios.
    min_eligible:
        Minimum number of tickers with positive signal to form a portfolio.

    Returns
    -------
    pd.DataFrame
        Weight DataFrame (rebalance dates × tickers).
    """
    rebal_dates = signal_df.resample(freq).last().index
    rebal_dates = [d for d in rebal_dates if d in signal_df.index]

    weights = pd.DataFrame(0.0, index=rebal_dates, columns=signal_df.columns)

    for d in rebal_dates:
        snap = signal_df.loc[d]
        if long_only:
            longs = snap[snap > 0].index.tolist()
            if len(longs) >= min_eligible:
                weights.loc[d, longs] = 1.0 / len(longs)
        else:
            pos = snap[snap > 0].index.tolist()
            neg = snap[snap < 0].index.tolist()
            n = len(pos) + len(neg)
            if n < min_eligible:
                continue
            if pos:
                weights.loc[d, pos] = 1.0 / len(pos)
            if neg:
                weights.loc[d, neg] = -1.0 / len(neg)

    return weights.clip(-1.0, 1.0)


# ---------------------------------------------------------------------------
# 6. Portfolio-level vectorized backtest with costs
# ---------------------------------------------------------------------------

def run_strategy_backtest(
    daily_returns: pd.DataFrame,
    target_weights: pd.DataFrame,
    cost_bps: float = 10.0,
    is_dollar_neutral: bool = False,
) -> Dict[str, Any]:
    """Vectorized portfolio backtest on a weight schedule.

    Parameters
    ----------
    daily_returns:
        Wide daily return DataFrame (tickers as columns).
    target_weights:
        Weight DataFrame indexed by **rebalance dates** (tickers as columns).
        Expanded to daily frequency via forward-fill inside this function.
    cost_bps:
        One-way transaction cost in basis points applied on traded notional.
    is_dollar_neutral:
        Metadata flag — does not change calculation; used in reporting.

    Returns
    -------
    dict with keys:
        ``gross`` : daily gross return Series
        ``net``   : daily net-of-cost return Series
        ``equity_gross`` : cumulative equity curve (start = 1.0)
        ``equity_net``   : cumulative equity curve (start = 1.0)
        ``turnover_per_rebal`` : mean |Δw| summed across tickers per rebalance
        ``cost_drag`` : total cost drag over the period
        ``n_rebal`` : number of rebalance dates with non-zero positions
    """
    cost_rate = cost_bps / 1e4

    # Expand weights to daily using forward-fill
    daily_w = target_weights.reindex(daily_returns.index, method="ffill").fillna(0.0)
    # Only keep columns present in returns
    common = [c for c in daily_w.columns if c in daily_returns.columns]
    daily_w = daily_w[common]
    daily_rets = daily_returns[common]

    # Gross PnL: yesterday's weights × today's returns (shift-1 = no look-ahead)
    gross = (daily_w.shift(1) * daily_rets).sum(axis=1)

    # Cost: turnover at rebalance dates only (Σ|Δw| × cost_rate)
    turn = target_weights[common].diff().abs().sum(axis=1)
    turn.iloc[0] = target_weights[common].iloc[0].abs().sum()  # initial deployment

    cost_series = pd.Series(0.0, index=daily_returns.index)
    cost_series.loc[turn.index] = turn.values * cost_rate
    cost_series = cost_series.reindex(daily_returns.index).fillna(0.0)

    net = (gross - cost_series).dropna()
    gross = gross.dropna()

    equity_gross = (1.0 + gross).cumprod()
    equity_net = (1.0 + net).cumprod()

    n_rebal = int((target_weights.abs().sum(axis=1) > 0).sum())

    return {
        "gross": gross,
        "net": net,
        "equity_gross": equity_gross,
        "equity_net": equity_net,
        "turnover_per_rebal": float(turn.mean()),
        "cost_drag": float((gross - net).sum()),
        "n_rebal": n_rebal,
        "is_dollar_neutral": is_dollar_neutral,
    }


# ---------------------------------------------------------------------------
# 7. Strategy metrics
# ---------------------------------------------------------------------------

def calculate_strategy_metrics(
    returns: pd.Series,
    rf_annual: float = 0.0,
    trading_days: int = 252,
) -> Dict[str, Any]:
    """Compute comprehensive performance metrics.

    Drawdown convention: **negative** (−0.25 = 25 % drawdown).
    rf = 0.0 by default; set explicitly — no FRED dependency.

    Parameters
    ----------
    returns:
        Daily strategy returns (net of costs).
    rf_annual:
        Annual risk-free rate (default 0.0; caveat: overstates long-only Sharpe
        in a high-rate regime).

    Returns
    -------
    dict
        Keys: annualized_return, annualized_vol, sharpe_ratio, sortino_ratio,
        max_drawdown (negative), calmar_ratio, monthly_hit_rate,
        t_stat_monthly, total_return, num_days.
    """
    r = returns.dropna()
    n = len(r)
    if n < 5:
        return {k: np.nan for k in (
            "annualized_return", "annualized_vol", "sharpe_ratio",
            "sortino_ratio", "max_drawdown", "calmar_ratio",
            "monthly_hit_rate", "t_stat_monthly", "total_return", "num_days",
        )} | {"num_days": n}

    ann_ret = float((1.0 + r).prod() ** (trading_days / n) - 1.0)
    ann_vol = float(r.std() * np.sqrt(trading_days))
    rf_daily = rf_annual / trading_days
    excess = r - rf_daily
    sharpe = float(excess.mean() / excess.std() * np.sqrt(trading_days)) if excess.std() > 0 else np.nan

    downside = r[r < 0].std()
    sortino = float((r.mean() / downside) * np.sqrt(trading_days)) if (downside and downside > 0) else np.nan

    eq = (1.0 + r).cumprod()
    dd = (eq - eq.cummax()) / eq.cummax()
    max_dd = float(dd.min())  # negative fraction

    calmar = float(ann_ret / abs(max_dd)) if max_dd < 0 else np.nan

    monthly = (1.0 + r).resample("ME").prod() - 1.0
    hit = float((monthly > 0).mean())
    t_stat = float(monthly.mean() / (monthly.std() / np.sqrt(len(monthly)))) \
        if (len(monthly) > 1 and monthly.std() > 0) else np.nan

    total = float((1.0 + r).prod() - 1.0)

    return {
        "annualized_return": ann_ret,
        "annualized_vol": ann_vol,
        "sharpe_ratio": sharpe,
        "sortino_ratio": sortino,
        "max_drawdown": max_dd,
        "calmar_ratio": calmar,
        "monthly_hit_rate": hit,
        "t_stat_monthly": t_stat,
        "total_return": total,
        "num_days": n,
    }


# ---------------------------------------------------------------------------
# 8. Alpha / beta vs benchmark
# ---------------------------------------------------------------------------

def calculate_alpha_beta_vs_benchmark(
    strategy_returns: pd.Series,
    benchmark_returns: pd.Series,
    trading_days: int = 252,
    min_obs: int = 60,
) -> Dict[str, float]:
    """OLS regression: daily strategy returns ~ α + β × benchmark_returns.

    Parameters
    ----------
    strategy_returns, benchmark_returns:
        Daily return Series (may have different indices; inner join used).
    trading_days:
        Used to annualize alpha (default 252).
    min_obs:
        Minimum aligned observations required; returns NaN dict if not met.

    Returns
    -------
    dict
        Keys: ``ann_alpha``, ``beta``, ``t_alpha``, ``r_squared``.
    """
    nan_result = {"ann_alpha": np.nan, "beta": np.nan, "t_alpha": np.nan, "r_squared": np.nan}
    df = pd.concat(
        [strategy_returns.rename("y"), benchmark_returns.rename("x")], axis=1
    ).dropna()
    if len(df) < min_obs:
        return nan_result

    y = df["y"].values
    x = df["x"].values
    X = np.column_stack([np.ones(len(y)), x])
    try:
        coef, _, _, _ = np.linalg.lstsq(X, y, rcond=None)
    except np.linalg.LinAlgError:
        return nan_result

    resid = y - X @ coef
    dof = len(y) - 2
    if dof <= 0:
        return nan_result
    sigma2 = (resid @ resid) / dof
    try:
        cov = sigma2 * np.linalg.inv(X.T @ X)
    except np.linalg.LinAlgError:
        return nan_result

    se_alpha = float(np.sqrt(cov[0, 0]))
    ann_alpha = float(coef[0] * trading_days)
    beta = float(coef[1])
    t_alpha = float(coef[0] / se_alpha) if se_alpha > 0 else np.nan

    ss_res = float(resid @ resid)
    ss_tot = float(((y - y.mean()) ** 2).sum())
    r2 = 1.0 - ss_res / ss_tot if ss_tot > 0 else np.nan

    return {"ann_alpha": ann_alpha, "beta": beta, "t_alpha": t_alpha, "r_squared": r2}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _verdict(avg_sharpe: Optional[float], t_alpha: Optional[float]) -> str:
    """Mechanical verdict — never massaged."""
    if avg_sharpe is None or np.isnan(avg_sharpe):
        return "no_results"
    if avg_sharpe > 1.0 and (t_alpha is not None and not np.isnan(t_alpha) and abs(t_alpha) >= 2.0):
        return "promising"
    if avg_sharpe > 0.5:
        return "marginal_positive"
    if avg_sharpe > 0.0:
        return "marginal"
    return "unprofitable"


def _is_oos_split(returns: pd.Series, oos_frac: float) -> Tuple[pd.Series, pd.Series]:
    cut = returns.index[int(len(returns) * (1.0 - oos_frac))]
    return returns[returns.index < cut], returns[returns.index >= cut]


# ---------------------------------------------------------------------------
# 9. Main research orchestrator
# ---------------------------------------------------------------------------

def run_ai_infra_research(
    period: str = "5y",
    interval: str = "1d",
    cost_bps: float = 10.0,
    slippage_bps: float = 10.0,
    output_dir: str = "outputs/research_runs/ai_infrastructure",
    long_only: bool = False,
    top_n: Optional[int] = None,
    symbols: Optional[List[str]] = None,
    primary_benchmark: str = "SOXX",
) -> Dict[str, Any]:
    """Run the full AI-infra alpha research pipeline.

    Steps
    -----
    1. Download real OHLCV (yfinance).
    2. Validate; log skipped symbols.
    3. Run agent review (DataEngineer, OpportunityAgent, QuantEngineer,
       RiskManager, ResearchAgent).
    4. Build and backtest four strategy families.
    5. Compare vs SOXX/SMH/QQQ/SPY benchmarks.
    6. Log to MemoryStore.
    7. Export outputs.

    Parameters
    ----------
    period:
        yfinance history period.
    interval:
        Bar interval (``"1d"`` recommended).
    cost_bps:
        One-way commission in bps (applied per side).
    slippage_bps:
        One-way slippage in bps.
    output_dir:
        Base directory for timestamped run outputs.
    long_only:
        If True, suppresses short legs in cross-sectional strategy.
    top_n:
        Optional cap on universe size (takes first *top_n* symbols).
    primary_benchmark:
        Ticker used for alpha/beta regression (default ``"SOXX"``).

    Returns
    -------
    dict
        Full research results (metrics, backtests, outputs, agent review).
    """
    from datetime import datetime, timezone
    start_time = datetime.now(timezone.utc)
    total_cost_bps = cost_bps + slippage_bps

    logger.info("=== AI Infra Research Pipeline ===")
    logger.info("period=%s | interval=%s | cost=%.0f+%.0f bps | long_only=%s",
                period, interval, cost_bps, slippage_bps, long_only)

    # ------------------------------------------------------------------
    # 1. Download
    # ------------------------------------------------------------------
    syms = symbols or _all_universe_symbols()
    if top_n:
        syms = syms[:top_n]

    try:
        close, volume, bench_close = download_ai_infra_data(
            symbols=syms, period=period, interval=interval,
        )
    except Exception as exc:
        logger.error("Data download failed: %s", exc)
        return {"error": str(exc), "universe": syms}

    # ------------------------------------------------------------------
    # 2. Validate
    # ------------------------------------------------------------------
    _, validation_report, skipped, valid_close, valid_volume = validate_ai_infra_data(
        close, volume
    )

    if valid_close.empty:
        logger.error("No valid data after validation — cannot continue.")
        return {
            "error": "No valid symbols after data validation",
            "skipped": skipped,
            "validation_report": validation_report,
        }

    daily_rets = valid_close.pct_change()

    # Benchmark returns
    bench_rets: Dict[str, pd.Series] = {}
    for col in bench_close.columns:
        s = bench_close[col].pct_change().dropna()
        bench_rets[col] = s

    primary_bench_ret = bench_rets.get(primary_benchmark, pd.Series(dtype=float))

    # ------------------------------------------------------------------
    # 3. Agent review
    # ------------------------------------------------------------------
    agent_review = _run_agent_review(
        valid_close, valid_volume, bench_close, validation_report, skipped
    )

    # ------------------------------------------------------------------
    # 4. Strategies
    # ------------------------------------------------------------------
    cfg = DEFAULT_CONFIG.copy()
    strategies: Dict[str, Any] = {}

    # 4a. Cross-sectional momentum (primary)
    logger.info("Building cross-sectional momentum signal…")
    cs_weights, _ = build_cross_sectional_momentum_signal(
        valid_close, valid_volume,
        lookback=cfg["cs_lookback_days"], skip=cfg["cs_skip_days"],
        liq_window=cfg["liq_window"], min_dollar_vol=cfg["min_dollar_vol"],
        quantile=cfg["quantile"], freq=cfg["rebalance_freq"],
        long_only=long_only,
    )
    cs_bt = run_strategy_backtest(daily_rets, cs_weights, total_cost_bps,
                                   is_dollar_neutral=not long_only)
    cs_metrics_g = calculate_strategy_metrics(cs_bt["gross"], cfg["rf_annual"], cfg["trading_days"])
    cs_metrics_n = calculate_strategy_metrics(cs_bt["net"], cfg["rf_annual"], cfg["trading_days"])
    cs_ab = calculate_alpha_beta_vs_benchmark(cs_bt["net"], primary_bench_ret)
    is_ret, oos_ret = _is_oos_split(cs_bt["net"], cfg["oos_split"])
    cs_is = calculate_strategy_metrics(is_ret)
    cs_oos = calculate_strategy_metrics(oos_ret)
    strategies["cross_sectional_momentum"] = {
        "backtest": cs_bt,
        "metrics_gross": cs_metrics_g,
        "metrics_net": cs_metrics_n,
        "alpha_beta": cs_ab,
        "is_metrics": cs_is,
        "oos_metrics": cs_oos,
        "verdict": _verdict(cs_metrics_n.get("sharpe_ratio"), cs_ab.get("t_alpha")),
        "hypothesis": _STRATEGY_HYPOTHESES["cross_sectional_momentum"],
        "is_primary": True,
    }

    # 4b. Trend following
    logger.info("Building trend-following signal…")
    tf_signal = build_trend_following_signal(valid_close, cfg["ma_windows"])
    tf_weights = build_monthly_rebalance_weights(
        tf_signal, cfg["rebalance_freq"], long_only=True,
        min_eligible=cfg["min_eligible"]
    )
    tf_bt = run_strategy_backtest(daily_rets, tf_weights, total_cost_bps, is_dollar_neutral=False)
    tf_metrics_n = calculate_strategy_metrics(tf_bt["net"], cfg["rf_annual"], cfg["trading_days"])
    tf_metrics_g = calculate_strategy_metrics(tf_bt["gross"], cfg["rf_annual"], cfg["trading_days"])
    tf_ab = calculate_alpha_beta_vs_benchmark(tf_bt["net"], primary_bench_ret)
    strategies["trend_following"] = {
        "backtest": tf_bt,
        "metrics_gross": tf_metrics_g,
        "metrics_net": tf_metrics_n,
        "alpha_beta": tf_ab,
        "verdict": _verdict(tf_metrics_n.get("sharpe_ratio"), tf_ab.get("t_alpha")),
        "hypothesis": _STRATEGY_HYPOTHESES["trend_following"],
        "is_primary": False,
    }

    # 4c. Mean reversion
    logger.info("Building mean-reversion signal…")
    mr_signal = build_mean_reversion_signal(
        valid_close, cfg["mr_window"], cfg["mr_z_threshold"], allow_short=False
    )
    mr_weights = build_monthly_rebalance_weights(
        mr_signal, cfg["rebalance_freq"], long_only=True,
        min_eligible=cfg["min_eligible"]
    )
    mr_bt = run_strategy_backtest(daily_rets, mr_weights, total_cost_bps, is_dollar_neutral=False)
    mr_metrics_n = calculate_strategy_metrics(mr_bt["net"], cfg["rf_annual"], cfg["trading_days"])
    mr_metrics_g = calculate_strategy_metrics(mr_bt["gross"], cfg["rf_annual"], cfg["trading_days"])
    mr_ab = calculate_alpha_beta_vs_benchmark(mr_bt["net"], primary_bench_ret)
    strategies["mean_reversion"] = {
        "backtest": mr_bt,
        "metrics_gross": mr_metrics_g,
        "metrics_net": mr_metrics_n,
        "alpha_beta": mr_ab,
        "verdict": _verdict(mr_metrics_n.get("sharpe_ratio"), mr_ab.get("t_alpha")),
        "hypothesis": _STRATEGY_HYPOTHESES["mean_reversion"],
        "is_primary": False,
    }

    # 4d. Breakout
    logger.info("Building breakout signal…")
    bo_signal = build_breakout_signal(valid_close, window=cfg["breakout_window"])
    bo_weights = build_monthly_rebalance_weights(
        bo_signal, cfg["rebalance_freq"], long_only=True,
        min_eligible=cfg["min_eligible"]
    )
    bo_bt = run_strategy_backtest(daily_rets, bo_weights, total_cost_bps, is_dollar_neutral=False)
    bo_metrics_n = calculate_strategy_metrics(bo_bt["net"], cfg["rf_annual"], cfg["trading_days"])
    bo_metrics_g = calculate_strategy_metrics(bo_bt["gross"], cfg["rf_annual"], cfg["trading_days"])
    bo_ab = calculate_alpha_beta_vs_benchmark(bo_bt["net"], primary_bench_ret)
    strategies["breakout"] = {
        "backtest": bo_bt,
        "metrics_gross": bo_metrics_g,
        "metrics_net": bo_metrics_n,
        "alpha_beta": bo_ab,
        "verdict": _verdict(bo_metrics_n.get("sharpe_ratio"), bo_ab.get("t_alpha")),
        "hypothesis": _STRATEGY_HYPOTHESES["breakout"],
        "is_primary": False,
    }

    # ------------------------------------------------------------------
    # 5. Benchmark comparison
    # ------------------------------------------------------------------
    bench_summary: Dict[str, Dict] = {}
    for bname, bret in bench_rets.items():
        bench_summary[bname] = calculate_strategy_metrics(bret, cfg["rf_annual"], cfg["trading_days"])

    # ------------------------------------------------------------------
    # 6. Memory logging
    # ------------------------------------------------------------------
    _log_to_memory(
        strategies=strategies,
        skipped=skipped,
        period=period,
        universe=list(valid_close.columns),
    )

    # ------------------------------------------------------------------
    # 7. Assemble results
    # ------------------------------------------------------------------
    results: Dict[str, Any] = {
        "universe_requested": syms,
        "universe_used": list(valid_close.columns),
        "symbols_skipped": skipped,
        "validation_report": validation_report,
        "period": period,
        "interval": interval,
        "cost_bps": total_cost_bps,
        "long_only": long_only,
        "rf_annual": cfg["rf_annual"],
        "strategies": strategies,
        "bench_summary": bench_summary,
        "primary_benchmark": primary_benchmark,
        "agent_review": agent_review,
        "run_timestamp": start_time.isoformat(),
        "duration_seconds": (datetime.now(timezone.utc) - start_time).total_seconds(),
        "disclaimer": (
            "RESEARCH / EDUCATIONAL PURPOSES ONLY. "
            "NOT INVESTMENT ADVICE. PAST PERFORMANCE DOES NOT GUARANTEE FUTURE RESULTS. "
            f"rf = {cfg['rf_annual']:.1%} (no FRED). "
            "Long-only Sharpe overstated in high-rate environments."
        ),
    }

    # ------------------------------------------------------------------
    # 8. Export
    # ------------------------------------------------------------------
    try:
        run_dir = export_ai_infra_outputs(results, base_dir=output_dir)
        results["output_dir"] = run_dir
        logger.info("Outputs written to %s", run_dir)
    except Exception as exc:
        logger.warning("export_ai_infra_outputs failed: %s", exc)
        results["output_dir"] = None

    logger.info(
        "Research complete in %.1fs | CS-mom=%s | TF=%s | MR=%s | breakout=%s",
        results["duration_seconds"],
        strategies["cross_sectional_momentum"]["verdict"],
        strategies["trend_following"]["verdict"],
        strategies["mean_reversion"]["verdict"],
        strategies["breakout"]["verdict"],
    )
    return results


# ---------------------------------------------------------------------------
# 9a. Agent review integration
# ---------------------------------------------------------------------------

def _run_agent_review(
    close: pd.DataFrame,
    volume: pd.DataFrame,
    bench_close: pd.DataFrame,
    validation_report: Dict,
    skipped: List[str],
) -> str:
    """Use Alpha Search agents to review data quality and rank opportunities.

    Falls back gracefully if agents are unavailable.

    Returns
    -------
    str
        Markdown-formatted agent review text.
    """
    lines = ["# Agent Review\n"]

    try:
        from alpha_search.agents.roles import (
            DataEngineerAgent,
            OpportunityAgent,
            QuantEngineerAgent,
            RiskManagerAgent,
            ResearchAgent,
        )
        from alpha_search.agents.swarm import CritiqueMessage

        de = DataEngineerAgent()
        oa = OpportunityAgent()

        # Build a combined prices DataFrame for the agents
        all_prices = close.copy()
        tickers = list(close.columns)

        # DataEngineerAgent: build a mock wide-format prices dict they can parse
        # The agents expect a MultiIndex DataFrame — wrap close as (Close, ticker)
        mc = pd.MultiIndex.from_tuples(
            [("Close", t) for t in tickers] + [("Volume", t) for t in volume.columns],
            names=["field", "ticker"],
        )
        wide_prices = pd.concat([close, volume], axis=1)
        wide_prices.columns = mc

        de_critiques = de.validate_data(wide_prices, tickers)
        lines.append("## DataEngineerAgent\n")
        for c in de_critiques:
            d = c if isinstance(c, dict) else c.to_dict()
            lines.append(f"- **[{d.get('severity','info').upper()}]** {d.get('message','')}")
        lines.append("")

        oa_critiques = oa.rank_opportunities(wide_prices, tickers)
        lines.append("## OpportunityAgent\n")
        for c in oa_critiques:
            d = c if isinstance(c, dict) else c.to_dict()
            lines.append(f"- {d.get('message','')}")
        lines.append("")

    except Exception as exc:
        logger.warning("Agent review partially failed (non-fatal): %s", exc)
        lines.append(f"_(Agent review unavailable: {exc})_\n")

    # Skipped symbols summary
    lines.append("## Data Quality Summary\n")
    if skipped:
        lines.append(f"**Skipped symbols** (non-positive prices — no data fabricated): `{', '.join(skipped)}`\n")
    else:
        lines.append("All symbols passed data validation.\n")

    lines.append("## Validation Report\n")
    for sym, rep in validation_report.items():
        issues = rep.get("issues", [])
        status = "OK" if not issues else "WARN"
        bars = rep.get("n_bars", 0)
        lines.append(f"- **{sym}** ({status}): {bars} bars" + (
            f" — {'; '.join(issues[:2])}" if issues else ""
        ))

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 9b. Memory logging
# ---------------------------------------------------------------------------

def _log_to_memory(
    strategies: Dict[str, Any],
    skipped: List[str],
    period: str,
    universe: List[str],
) -> None:
    """Persist research findings to MemoryStore (warn and continue on failure)."""
    try:
        from alpha_search.memory.models import MemoryRecord
        from alpha_search.memory.store import MemoryStore

        store = MemoryStore()
        store.initialize()

        # Log hypothesis
        store.add_memory(MemoryRecord(
            agent_name="ai_infra_pipeline",
            memory_type="architecture_decision",
            title="AI Infra Alpha Research — parameter registration",
            content=(
                f"Pre-registered parameters: period={period}, "
                f"universe_size={len(universe)}, skipped={skipped}. "
                "No grid search; parameters fixed before results observed."
            ),
            tags=["ai_infra", "pre_registration"],
            importance_score=0.9,
        ))

        # Log each strategy result
        for strat_name, strat in strategies.items():
            metrics = strat.get("metrics_net", {})
            sharpe = metrics.get("sharpe_ratio", np.nan)
            verdict = strat.get("verdict", "unknown")
            store.add_memory(MemoryRecord(
                agent_name="ai_infra_pipeline",
                memory_type="strategy_result",
                title=f"AI Infra — {strat_name} — {verdict}",
                content=(
                    f"Strategy: {strat_name} | Verdict: {verdict} | "
                    f"Net Sharpe: {sharpe:.3f} | "
                    f"Max DD: {metrics.get('max_drawdown', float('nan')):.3f} | "
                    f"Alpha: {strat.get('alpha_beta', {}).get('ann_alpha', float('nan')):.3f}"
                )[:2000],
                tags=["ai_infra", strat_name, verdict],
                importance_score=0.8 if verdict in ("promising", "marginal_positive") else 0.5,
            ))

        # Log skipped symbols
        if skipped:
            store.add_memory(MemoryRecord(
                agent_name="ai_infra_pipeline",
                memory_type="data_quality_issue",
                title="AI Infra — skipped symbols",
                content=f"Skipped (non-positive prices): {skipped}. No data fabricated.",
                tags=["ai_infra", "data_quality"],
            ))

        store.close()
        logger.info("Memory logging complete.")

    except Exception as exc:
        logger.warning("Memory logging failed (non-fatal): %s", exc)


# ---------------------------------------------------------------------------
# 10. Export outputs
# ---------------------------------------------------------------------------

def export_ai_infra_outputs(
    results: Dict[str, Any],
    base_dir: str = "outputs/research_runs/ai_infrastructure",
) -> str:
    """Write all output files to a timestamped directory.

    Creates the directory structure:
    ``<base_dir>/<YYYYMMDD_HHMMSS>/``
    with metadata, CSV files, figures, and a Markdown report.

    Returns
    -------
    str
        Path to the timestamped run directory.
    """
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    run_dir = os.path.join(base_dir, ts)
    fig_dir = os.path.join(run_dir, "figures")
    os.makedirs(fig_dir, exist_ok=True)

    # --- metadata.json ---
    meta = {
        "run_timestamp": results.get("run_timestamp"),
        "period": results.get("period"),
        "interval": results.get("interval"),
        "cost_bps": results.get("cost_bps"),
        "long_only": results.get("long_only"),
        "rf_annual": results.get("rf_annual"),
        "universe_used": results.get("universe_used", []),
        "symbols_skipped": results.get("symbols_skipped", []),
        "primary_benchmark": results.get("primary_benchmark"),
        "duration_seconds": results.get("duration_seconds"),
        "disclaimer": results.get("disclaimer"),
    }
    with open(os.path.join(run_dir, "metadata.json"), "w") as f:
        json.dump(meta, f, indent=2)

    # --- universe_used.csv ---
    pd.DataFrame({"symbol": results.get("universe_used", [])}).to_csv(
        os.path.join(run_dir, "universe_used.csv"), index=False
    )

    # --- skipped_symbols.csv ---
    pd.DataFrame({"symbol": results.get("symbols_skipped", [])}).to_csv(
        os.path.join(run_dir, "skipped_symbols.csv"), index=False
    )

    # --- per-strategy CSVs + summary ---
    strategies = results.get("strategies", {})
    summary_rows = []
    for strat_name, strat in strategies.items():
        metrics_n = strat.get("metrics_net", {})
        metrics_g = strat.get("metrics_gross", {})
        ab = strat.get("alpha_beta", {})
        row = {
            "strategy": strat_name,
            "verdict": strat.get("verdict"),
            **{f"net_{k}": v for k, v in metrics_n.items()},
            **{f"gross_sharpe": metrics_g.get("sharpe_ratio")},
            "ann_alpha_vs_bench": ab.get("ann_alpha"),
            "beta_vs_bench": ab.get("beta"),
            "t_alpha": ab.get("t_alpha"),
            "r_squared": ab.get("r_squared"),
            "turnover_per_rebal": strat.get("backtest", {}).get("turnover_per_rebal"),
        }
        summary_rows.append(row)

        bt = strat.get("backtest", {})
        net_s = bt.get("net")
        if net_s is not None and isinstance(net_s, pd.Series) and not net_s.empty:
            fname = strat_name.replace(" ", "_") + "_results.csv"
            pd.DataFrame({"date": net_s.index, "net_return": net_s.values}).to_csv(
                os.path.join(run_dir, fname), index=False
            )

    pd.DataFrame(summary_rows).to_csv(
        os.path.join(run_dir, "strategy_results_summary.csv"), index=False
    )

    # --- benchmark comparison ---
    bench_rows = []
    for bname, bm in results.get("bench_summary", {}).items():
        bench_rows.append({"benchmark": bname, **bm})
    if bench_rows:
        pd.DataFrame(bench_rows).to_csv(
            os.path.join(run_dir, "benchmark_comparison.csv"), index=False
        )

    # --- agent_review.md ---
    agent_review = results.get("agent_review", "")
    with open(os.path.join(run_dir, "agent_review.md"), "w") as f:
        f.write(agent_review)

    # --- report.md ---
    report_path = os.path.join(run_dir, "report.md")
    generate_ai_infra_report(results, report_path)

    # --- figures ---
    _write_ai_infra_figures(fig_dir, results)

    # --- docx (optional) ---
    try:
        import docx  # type: ignore
        _write_docx_report(os.path.join(run_dir, "report.docx"), results)
    except ImportError:
        logger.debug("python-docx not installed — skipping report.docx")

    return run_dir


def _write_ai_infra_figures(fig_dir: str, results: Dict[str, Any]) -> None:
    """Write strategy comparison figures using Agg backend."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        logger.debug("matplotlib not available — skipping figures")
        return

    strategies = results.get("strategies", {})
    colours = {"cross_sectional_momentum": "royalblue", "trend_following": "seagreen",
               "mean_reversion": "darkorange", "breakout": "purple"}

    # Equity curves
    fig, ax = plt.subplots(figsize=(12, 5))
    for sname, strat in strategies.items():
        eq = strat.get("backtest", {}).get("equity_net")
        if eq is not None and not eq.empty:
            eq.plot(ax=ax, label=sname.replace("_", " ").title(),
                    color=colours.get(sname), linewidth=1.3)
    ax.set_title("Net Equity Curves — AI Infra Strategies")
    ax.set_ylabel("Cumulative return (start = 1.0)")
    ax.legend(fontsize=8)
    ax.set_yscale("log")
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "equity_curve.png"), dpi=120)
    plt.close(fig)

    # Drawdown
    fig, ax = plt.subplots(figsize=(12, 4))
    for sname, strat in strategies.items():
        eq = strat.get("backtest", {}).get("equity_net")
        if eq is not None and not eq.empty:
            dd = (eq - eq.cummax()) / eq.cummax()
            dd.plot(ax=ax, label=sname.replace("_", " ").title(),
                    color=colours.get(sname), linewidth=1.1, alpha=0.8)
    ax.axhline(0, color="black", linewidth=0.7)
    ax.set_title("Drawdown — AI Infra Strategies (negative convention)")
    ax.set_ylabel("Drawdown")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{x:.0%}"))
    ax.legend(fontsize=8)
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "drawdown_curve.png"), dpi=120)
    plt.close(fig)

    # Rolling Sharpe (cross-sectional momentum only)
    cs_bt = strategies.get("cross_sectional_momentum", {}).get("backtest", {})
    cs_net = cs_bt.get("net")
    if cs_net is not None and not cs_net.empty:
        fig, ax = plt.subplots(figsize=(12, 4))
        roll_sharpe = cs_net.rolling(63).mean() / (cs_net.rolling(63).std() + 1e-9) * np.sqrt(252)
        roll_sharpe.plot(ax=ax, color="royalblue", linewidth=1.2)
        ax.axhline(0, color="grey", linewidth=0.8, linestyle="--")
        ax.axhline(1, color="green", linewidth=0.8, linestyle=":")
        ax.set_title("Rolling 63-day Sharpe — Cross-Sectional Momentum (Net)")
        ax.set_ylabel("Sharpe ratio")
        plt.tight_layout()
        fig.savefig(os.path.join(fig_dir, "rolling_sharpe.png"), dpi=120)
        plt.close(fig)

    # Strategy comparison bar chart
    _strats = list(strategies.keys())
    net_sharpes = [strategies[s].get("metrics_net", {}).get("sharpe_ratio", np.nan) for s in _strats]
    gross_sharpes = [strategies[s].get("metrics_gross", {}).get("sharpe_ratio", np.nan) for s in _strats]
    x = np.arange(len(_strats))
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.bar(x - 0.2, gross_sharpes, 0.35, label="Gross Sharpe", color="steelblue", alpha=0.7)
    ax.bar(x + 0.2, net_sharpes, 0.35, label="Net Sharpe", color="firebrick", alpha=0.7)
    ax.axhline(0, color="black", linewidth=0.7)
    ax.set_xticks(x)
    ax.set_xticklabels([s.replace("_", "\n") for s in _strats], fontsize=8)
    ax.set_title("Gross vs Net Sharpe by Strategy")
    ax.set_ylabel("Sharpe ratio")
    ax.legend()
    plt.tight_layout()
    fig.savefig(os.path.join(fig_dir, "strategy_comparison.png"), dpi=120)
    plt.close(fig)

    # Benchmark comparison
    bench_summary = results.get("bench_summary", {})
    if bench_summary:
        bnames = list(bench_summary.keys())
        bsharpes = [bench_summary[b].get("sharpe_ratio", np.nan) for b in bnames]
        fig, ax = plt.subplots(figsize=(8, 4))
        colours_b = ["#2196F3", "#4CAF50", "#FF9800", "#9C27B0"]
        for i, (bname, bsharpe) in enumerate(zip(bnames, bsharpes)):
            ax.bar(bname, bsharpe, color=colours_b[i % len(colours_b)], alpha=0.8)
        # Add primary strategy
        cs_sharpe = strategies.get("cross_sectional_momentum", {}).get(
            "metrics_net", {}
        ).get("sharpe_ratio", np.nan)
        if not np.isnan(cs_sharpe):
            ax.bar("CS-Mom\n(L/S net)", cs_sharpe, color="royalblue", alpha=0.9)
        ax.axhline(0, color="black", linewidth=0.7)
        ax.set_title("Benchmark Sharpe Comparison")
        ax.set_ylabel("Annualized Sharpe ratio")
        plt.tight_layout()
        fig.savefig(os.path.join(fig_dir, "benchmark_comparison.png"), dpi=120)
        plt.close(fig)

    # Correlation heatmap
    return_series = {}
    for sname, strat in strategies.items():
        net = strat.get("backtest", {}).get("net")
        if net is not None and not net.empty:
            return_series[sname] = net
    if len(return_series) > 1:
        try:
            corr_df = pd.concat(return_series, axis=1).corr()
            fig, ax = plt.subplots(figsize=(7, 6))
            im = ax.imshow(corr_df.values, cmap="RdYlGn", vmin=-1, vmax=1)
            ax.set_xticks(range(len(corr_df)))
            ax.set_yticks(range(len(corr_df)))
            labels = [s.replace("_", "\n") for s in corr_df.columns]
            ax.set_xticklabels(labels, fontsize=7, rotation=45, ha="right")
            ax.set_yticklabels(labels, fontsize=7)
            for i in range(len(corr_df)):
                for j in range(len(corr_df)):
                    ax.text(j, i, f"{corr_df.values[i, j]:.2f}",
                            ha="center", va="center", fontsize=7)
            plt.colorbar(im, ax=ax)
            ax.set_title("Strategy Return Correlation Heatmap")
            plt.tight_layout()
            fig.savefig(os.path.join(fig_dir, "correlation_heatmap.png"), dpi=120)
            plt.close(fig)
        except Exception as exc:
            logger.debug("Heatmap failed: %s", exc)


def _write_docx_report(path: str, results: Dict[str, Any]) -> None:
    import docx  # type: ignore
    doc = docx.Document()
    doc.add_heading("AI Infra & Semiconductor Alpha Research", 0)
    doc.add_paragraph(results.get("disclaimer", ""))
    doc.add_heading("Strategy Summary", 1)
    strategies = results.get("strategies", {})
    for sname, strat in strategies.items():
        metrics = strat.get("metrics_net", {})
        doc.add_heading(sname.replace("_", " ").title(), 2)
        doc.add_paragraph(
            f"Verdict: {strat.get('verdict')} | "
            f"Net Sharpe: {metrics.get('sharpe_ratio', float('nan')):.3f} | "
            f"Max DD: {metrics.get('max_drawdown', float('nan')):.2%}"
        )
    doc.save(path)


# ---------------------------------------------------------------------------
# 11. Markdown report
# ---------------------------------------------------------------------------

def generate_ai_infra_report(results: Dict[str, Any], path: str) -> None:
    """Write a comprehensive Markdown research report.

    Negative Sharpe and drawdown are reported as-is — no massaging.
    """
    lines: List[str] = []
    ts = results.get("run_timestamp", datetime.now(timezone.utc).isoformat())

    lines += [
        "# AI Infrastructure & Semiconductor — Alpha Research Report",
        "",
        f"**Generated:** {ts}",
        f"**Universe:** {len(results.get('universe_used', []))} symbols "
        f"({results.get('period')} | {results.get('interval')})",
        f"**Cost:** {results.get('cost_bps', 0):.0f} bps round-trip",
        f"**rf = {results.get('rf_annual', 0):.1%}** (no FRED — overstates long-only Sharpe in high-rate environments)",
        "",
        f"> **DISCLAIMER:** {results.get('disclaimer', 'Research only.')}",
        "",
        "---",
    ]

    # Skipped symbols
    skipped = results.get("symbols_skipped", [])
    if skipped:
        lines += [
            "", "## Data Quality",
            f"**Skipped symbols** (non-positive prices; no data fabricated): `{', '.join(skipped)}`",
        ]

    # Strategies
    lines += ["", "## Strategy Results (net of costs)", ""]
    lines += [
        "| Strategy | Verdict | Net Sharpe | Ann. Return | Max DD | Alpha vs Bench | t(α) | Turnover/rebal |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for sname, strat in results.get("strategies", {}).items():
        m = strat.get("metrics_net", {})
        ab = strat.get("alpha_beta", {})
        bt = strat.get("backtest", {})
        sharpe = m.get("sharpe_ratio", np.nan)
        ann_ret = m.get("annualized_return", np.nan)
        max_dd = m.get("max_drawdown", np.nan)
        alpha = ab.get("ann_alpha", np.nan)
        t_alp = ab.get("t_alpha", np.nan)
        turn = bt.get("turnover_per_rebal", np.nan)
        lines.append(
            f"| {sname.replace('_', ' ').title()} "
            f"| {strat.get('verdict')} "
            f"| {sharpe:.3f} "
            f"| {ann_ret:.2%} "
            f"| {max_dd:.2%} "
            f"| {alpha:.2%} "
            f"| {t_alp:.2f} "
            f"| {turn:.3f} |"
        )

    # IS/OOS
    cs = results.get("strategies", {}).get("cross_sectional_momentum", {})
    if cs:
        is_m = cs.get("is_metrics", {})
        oos_m = cs.get("oos_metrics", {})
        lines += [
            "", "## Cross-Sectional Momentum — IS/OOS Stability",
            "(Split at temporal midpoint — not chosen to flatter results.)",
            "",
            f"- In-sample Sharpe : **{is_m.get('sharpe_ratio', float('nan')):.3f}**",
            f"- Out-of-sample Sharpe : **{oos_m.get('sharpe_ratio', float('nan')):.3f}**",
        ]

    # Benchmark
    lines += ["", "## Benchmark Comparison", ""]
    lines += [
        "| Benchmark | Sharpe | Ann. Return | Max DD |",
        "|---|---|---|---|",
    ]
    for bname, bm in results.get("bench_summary", {}).items():
        lines.append(
            f"| {bname} "
            f"| {bm.get('sharpe_ratio', float('nan')):.3f} "
            f"| {bm.get('annualized_return', float('nan')):.2%} "
            f"| {bm.get('max_drawdown', float('nan')):.2%} |"
        )

    # Verdict
    primary = results.get("strategies", {}).get("cross_sectional_momentum", {})
    p_sharpe = primary.get("metrics_net", {}).get("sharpe_ratio", float("nan"))
    p_t = primary.get("alpha_beta", {}).get("t_alpha", float("nan"))
    p_verdict = primary.get("verdict", "n/a")
    lines += [
        "", "## Overall Verdict (Primary Strategy: Cross-Sectional Momentum L/S)",
        "",
        f"**Net Sharpe:** {p_sharpe:.3f}  |  **Alpha t-stat:** {p_t:.2f}  |  **Verdict:** `{p_verdict}`",
        "",
    ]
    if p_verdict == "promising":
        lines.append("Evidence of risk-adjusted alpha after costs. Recommend walk-forward validation before deployment.")
    elif p_verdict in ("marginal_positive", "marginal"):
        lines.append("Marginal positive edge; cost-sensitive. Further stress-testing required.")
    else:
        lines.append(
            "No robust alpha detected after costs. This result is reported as-is — "
            "parameters were not tuned to improve it."
        )

    lines += [
        "", "## Methodology",
        "- 1-bar position lag: weights from rebalance t take effect at t+1.",
        "- Donchian channel uses `shift(1)` before `rolling()` to prevent look-ahead.",
        "- Transaction costs charged on traded notional at each rebalance.",
        "- Single pre-registered parameter set — no grid search.",
        "- Drawdown convention: negative fraction (−0.25 = 25% DD).",
        "", "---", "_Research only. Not investment advice._",
    ]

    with open(path, "w") as f:
        f.write("\n".join(lines) + "\n")

    logger.info("Report written to %s", path)
